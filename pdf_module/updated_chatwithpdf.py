import os
import uuid
import pdfplumber
from docx import Document as DocxDocument  # ✅ New import for .docx
from dotenv import load_dotenv
from concurrent.futures import ThreadPoolExecutor
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
import google.generativeai as genai
from langchain_google_genai import ChatGoogleGenerativeAI
from pinecone import Pinecone, ServerlessSpec

# Load environment variables
load_dotenv()

PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# Configure APIs
pc = Pinecone(api_key=PINECONE_API_KEY)
genai.configure(api_key=GOOGLE_API_KEY)
embedding_model = genai.embed_content


# -------------------- Load and Chunk Documents --------------------
def load_and_chunk_documents(file_paths, chunk_size=1000, chunk_overlap=100):
    documents = []

    for file_path in file_paths:
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        metadata = {"page": i + 1, "source": file_path}
                        documents.append(Document(page_content=page_text, metadata=metadata))

        elif ext == ".docx":  # ✅ New Word document handling
            doc = DocxDocument(file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            if full_text:
                metadata = {"source": file_path}
                documents.append(Document(page_content=full_text, metadata=metadata))

        else:
            print(f"❌ Unsupported file type: {file_path}")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ".", " ", ""]
    )

    chunks = splitter.split_documents(documents)
    return chunks


# -------------------- Remaining Code (unchanged) --------------------
def get_gemini_embedding(text, model="models/text-embedding-004", task_type="retrieval_document"):
    try:
        res = embedding_model(model=model, content=text, task_type=task_type, title="chunk")
        return res["embedding"]
    except Exception as e:
        print(f"Embedding failed for a chunk: {e}")
        return None


def get_gemini_embeddings(texts, model="models/text-embedding-004", task_type="retrieval_document", max_workers=4):
    embeddings = []
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = [executor.submit(get_gemini_embedding, txt, model, task_type) for txt in texts]
        for future in futures:
            embeddings.append(future.result())
    return embeddings


def store_in_pinecone(chunks, embeddings, namespace="default", index_name="conversational-bi-index-2", dimension=768, batch_size=100):
    if index_name not in pc.list_indexes().names():
        print(f"Index '{index_name}' not found. Creating new index...")
        pc.create_index(
            name=index_name,
            dimension=dimension,
            metric="cosine",
            spec=ServerlessSpec(cloud="aws", region="us-east-1")
        )
        print(f"✅ Index '{index_name}' created successfully!")

    index = pc.Index(index_name)
    to_upsert = []
    for doc, embed in zip(chunks, embeddings):
        if embed is not None:
            vector = {
                "id": str(uuid.uuid4()),
                "values": embed,
                "metadata": {"text": doc.page_content, "page": doc.metadata.get("page", 0)}
            }
            to_upsert.append(vector)
            if len(to_upsert) >= batch_size:
                index.upsert(vectors=to_upsert, namespace=namespace)
                print(f"✅ Stored {len(to_upsert)} vectors in index '{index_name}' under namespace '{namespace}'")
                to_upsert = []

    if to_upsert:
        index.upsert(vectors=to_upsert, namespace=namespace)
        print(f"✅ Stored {len(to_upsert)} vectors in index '{index_name}' under namespace '{namespace}'")


def retrieve_from_pinecone(query, top_k=6, namespace="default", index_name="llm-chatbot"):
    index = pc.Index(index_name)
    query_embedding = get_gemini_embeddings([query])[0]
    results = index.query(vector=query_embedding, top_k=top_k, include_metadata=True, namespace=namespace)
    return [match['metadata']['text'] for match in results['matches']]


def generate_response(question, context_chunks):
    context = "\n\n".join(context_chunks)
    prompt_template = """
    You are a knowledgeable AI assistant. Answer the following question in detail, providing thorough explanations, examples, and key insights.

Context: {context}

Question: {question}
If user's question is type of greetings like ("hi","how are you","what you can do"),then answer this questions like ("Hi there! 👋 I’m your document assistant. You can upload PDFs or Word files and ask me anything about their content.")
If the user's question is general (like "what is this document about", 
"tell me about document", or similar), then **summarize** the main ideas 
of the document in 4–6 sentences.
Provide a well-structured and comprehensive response in markdown format. Use bullet points, headings, and code blocks only where appropriate. 
Base your answer strictly on the provided CONTEXT and if the answer is not found in the CONTEXT, 
reply with "I couldn't find the answer in the provided document.
        """
    llm = ChatGoogleGenerativeAI(model="gemini-2.0-flash", temperature=0.2, max_output_tokens=512)
    return llm.invoke(prompt_template.format(context=context, question=question)).content


def run_rag_pipeline(query, namespace="default", index_name="llm-chatbot"):
    retrieved = retrieve_from_pinecone(query, namespace=namespace, index_name=index_name)
    print(f"Retrieved {len(retrieved)} chunks for query: '{query}'")
    if len(retrieved) > 0:
        print("Sample context snippet:", retrieved[0][:200])
    else:
        print("⚠️ No relevant chunks found in Pinecone for this query.")
    response=generate_response(query, retrieved)
    return response
